"""
Utility functions for XPS conversion
"""
import os
import fitz
from PIL import Image
from django.conf import settings
import uuid
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Inches, Cm


class XPSConverter:
    """Main converter class for XPS files"""

    # --- CHANGED: Added original_filename to the constructor ---
    def __init__(self, xps_file_path, original_filename):
        self.xps_file_path = xps_file_path
        self.original_filename = original_filename
        self.doc = None

    def __enter__(self):
        """Open XPS document"""
        try:
            self.doc = fitz.open(self.xps_file_path)
            return self
        except Exception as e:
            raise Exception(f"Failed to open XPS file: {str(e)}")

    def __exit__(self, exc_type, exc_val, exc_tb):
        """Close XPS document"""
        if self.doc:
            self.doc.close()

    def get_page_count(self):
        """Get total number of pages"""
        if self.doc:
            return len(self.doc)
        return 0

    def validate_page_count(self):
        """Validate if page count is within limit"""
        page_count = self.get_page_count()
        if page_count > settings.XPS_MAX_PAGES:
            raise Exception(
                f"XPS file has {page_count} pages. Maximum allowed is {settings.XPS_MAX_PAGES} pages"
            )
        return page_count

    def convert_to_images(self):
        # ... (unchanged - no filename needed here)
        try:
            page_count = self.validate_page_count()
            images = []
            
            conversion_id = str(uuid.uuid4())
            output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', conversion_id)
            os.makedirs(output_dir, exist_ok=True)
            
            for page_num in range(page_count):
                page = self.doc[page_num]
                
                mat = fitz.Matrix(settings.IMAGE_DPI / 72, settings.IMAGE_DPI / 72)
                pix = page.get_pixmap(matrix=mat)
                
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                image_filename = f"page_{page_num + 1}.jpg"
                image_path = os.path.join(output_dir, image_filename)
                img.save(image_path, "JPEG", quality=settings.IMAGE_QUALITY, optimize=True)
                
                image_url = f"/media/converted/{conversion_id}/{image_filename}"
                images.append(image_url)
            
            return {
                'success': True,
                'images': images,
                'total_pages': page_count,
                'conversion_id': conversion_id
            }
            
        except Exception as e:
            raise Exception(f"Image conversion failed: {str(e)}")

    def convert_to_pdf(self):
        try:
            page_count = self.validate_page_count()
            
            conversion_id = str(uuid.uuid4())
            output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', conversion_id)
            os.makedirs(output_dir, exist_ok=True)
            
            # --- CHANGED: Use original filename (with new extension) ---
            base_name = os.path.splitext(self.original_filename)[0]
            pdf_filename = f"{base_name}.pdf"
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            pdf_bytes = self.doc.convert_to_pdf()
            
            with open(pdf_path, "wb") as pdf_file:
                pdf_file.write(pdf_bytes)
                
            pdf_url = f"/media/converted/{conversion_id}/{pdf_filename}"
            
            return {
                'success': True,
                'pdf_url': pdf_url,
                'total_pages': page_count,
                'conversion_id': conversion_id
            }
            
        except Exception as e:
            raise Exception(f"PDF conversion failed: {str(e)}")


    def convert_to_docx_text_only(self):
        # ... (similar changes to convert_to_docx for output filename)
        try:
            page_count = self.validate_page_count()
            
            conversion_id = str(uuid.uuid4())
            output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', conversion_id)
            os.makedirs(output_dir, exist_ok=True)
            
            doc = Document()
            doc.add_heading('XPS Conversion (Text Only)', 0)
            
            for page_num in range(page_count):
                page = self.doc[page_num]
                text = page.get_text()
                
                doc.add_heading(f'Page {page_num + 1}', level=1)
                doc.add_paragraph(text)
                doc.add_page_break()
            
            # --- CHANGED: Use original filename (with new extension) ---
            base_name = os.path.splitext(self.original_filename)[0]
            docx_filename = f"{base_name}_text_only.docx"
            docx_path = os.path.join(output_dir, docx_filename)
            doc.save(docx_path)
            
            docx_url = f"/media/converted/{conversion_id}/{docx_filename}"
            
            return {
                'success': True,
                'docx_url': docx_url,
                'total_pages': page_count,
                'conversion_id': conversion_id
            }
            
        except Exception as e:
            raise Exception(f"DOCX (Text Only) conversion failed: {str(e)}")


    def convert_to_docx(self):
        try:
            image_conversion_result = self.convert_to_images()
            if not image_conversion_result['success']:
                raise Exception("Image conversion failed, cannot create image-based DOCX.")

            images_paths = [
                os.path.join(settings.MEDIA_ROOT, url.lstrip('/media/'))
                for url in image_conversion_result['images']
            ]
            
            conversion_id = image_conversion_result['conversion_id']
            output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', conversion_id)
            
            doc = Document()

            section = doc.sections[0]
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)
            
            total_pages = len(images_paths)
            
            max_width = Inches(7.5) 
            
            for i, img_path in enumerate(images_paths):
                doc.add_picture(img_path, width=max_width)
                
                # if i < total_pages - 1:
                #     doc.add_page_break()
                
            # --- CHANGED: Use original filename (with new extension) ---
            base_name = os.path.splitext(self.original_filename)[0]
            docx_filename = f"{base_name}_image_based.docx"
            docx_path = os.path.join(output_dir, docx_filename)
            doc.save(docx_path)
            
            docx_url = f"/media/converted/{conversion_id}/{docx_filename}"
            
            return {
                'success': True,
                'docx_url': docx_url,
                'total_pages': total_pages,
                'conversion_id': conversion_id
            }
            
        except Exception as e:
            raise Exception(f"DOCX (Image Based) conversion failed: {str(e)}")


    def extract_text(self):
        # ... (unchanged - no filename needed here)
        try:
            page_count = self.validate_page_count()
            text_content = []
            
            for page_num in range(page_count):
                page = self.doc[page_num]
                text = page.get_text()
                
                text_content.append({
                    'page': page_num + 1,
                    'text': text
                })
            
            return {
                'success': True,
                'text_content': text_content,
                'total_pages': page_count
            }
            
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")


# --- CHANGED: Updated function signature and logic ---
def save_uploaded_file(uploaded_file, original_filename):
    """Save uploaded file temporarily"""
    try:
        # Create unique filename (the temporary file still needs a UUID)
        file_id = str(uuid.uuid4())
        temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        # The temporary file is named with the UUID, not the original name
        file_path = os.path.join(temp_dir, f"{file_id}.xps")
        
        # Save file
        with open(file_path, 'wb+') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)
        
        return file_path
    except Exception as e:
        raise Exception(f"Failed to save uploaded file: {str(e)}")


def cleanup_file(file_path):
    # ... (unchanged)
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        print(f"Failed to delete file {file_path}: {str(e)}")


def cleanup_directory(directory_path):
    # ... (unchanged)
    try:
        if os.path.exists(directory_path):
            shutil.rmtree(directory_path)
    except Exception as e:
        print(f"Failed to delete directory {directory_path}: {str(e)}")